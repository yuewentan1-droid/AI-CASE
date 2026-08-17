"""导出引擎测试"""
import os

from com.aiase.services.export.exporter import ExportEngine
from com.aiase.config.settings import Config

CASES = [
    {'主模块': '登录', '子模块1': '账号', '子模块2': '', '测试标题': '正确账号登录', '测试步骤': '输入账号\n点击登录', '测试结果': '登录成功\n已退出登录'},
    {'主模块': '登录', '子模块1': '账号', '子模块2': '', '测试标题': '错误密码', '测试步骤': '输入错误密码\n点击登录', '测试结果': '提示密码错误'},
]
TEMPLATE = {'fields': ['主模块', '子模块1', '子模块2', '测试标题', '测试步骤', '测试结果']}
os.makedirs(Config.TEMP_PATH, exist_ok=True)


def _path(fmt):
    return os.path.join(Config.TEMP_PATH, f'test_export.{fmt}')


def test_export_json():
    p = ExportEngine().export(CASES, TEMPLATE, 'json', 't1')
    assert os.path.exists(p)


def test_export_excel():
    p = ExportEngine().export(CASES, TEMPLATE, 'excel', 't1')
    assert os.path.exists(p)
    # 含换行的步骤/结果单元格开启自动换行，保证逐行可见、一一对应
    from openpyxl import load_workbook
    ws = load_workbook(p).active
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            if cell.value and '\n' in str(cell.value):
                assert cell.alignment.wrap_text is True


def test_export_word():
    p = ExportEngine().export(CASES, TEMPLATE, 'word', 't1')
    assert os.path.exists(p)
    # 步骤/结果每行拆成独立段落，保证换行可见、一一对应
    from docx import Document
    table = Document(p).tables[0]
    # 字段顺序：主模块/子模块1/子模块2/测试标题/测试步骤/测试结果，测试步骤列索引为 4
    step_cell = table.rows[1].cells[4]
    assert [x.text for x in step_cell.paragraphs] == ['1. 输入账号', '2. 点击登录']


def test_export_xmind():
    p = ExportEngine().export(CASES, TEMPLATE, 'xmind', 't1')
    with open(p, encoding='utf-8') as f:
        content = f.read()
    assert '正确账号登录' in content
    # 步骤 N 下挂对应序号的结果 N，一一对应
    assert '<node TEXT="1. 输入账号">' in content
    assert '<node TEXT="1. 登录成功"/>' in content
    assert '<node TEXT="2. 点击登录">' in content
    assert '<node TEXT="2. 已退出登录"/>' in content


def test_export_html():
    p = ExportEngine().export(CASES, TEMPLATE, 'html', 't1')
    assert os.path.exists(p)
    with open(p, encoding='utf-8') as f:
        content = f.read()
    # 步骤/结果用 <br> 分隔，保证逐行显示、一一对应（而不是被折叠成一行）
    assert '<br>' in content
    assert '1. 输入账号<br>2. 点击登录' in content


def test_export_txt():
    p = ExportEngine().export(CASES, TEMPLATE, 'txt', 't1')
    assert os.path.exists(p)
