"""需求评审内容导出 - 输出 analysis（风险/影响/功能/摘要）+ review（评审点）到 html/excel/word/xmind(.mm)"""
import xml.sax.saxutils as sax

from com.aiase.services.export.base_exporter import BaseExporter


def _esc(text):
    return sax.escape(str(text or ''))


def _review_stats(review):
    """统计评审 issue 高/中/低数量"""
    counts = {'高': 0, '中': 0, '低': 0}
    for i in (review.get('issues') or []):
        lv = i.get('level', '')
        if lv in counts:
            counts[lv] += 1
    return counts


class AnalysisExporter(BaseExporter):
    """把需求评审的全部内容点（分析内容 + 用例评审点）导出为指定格式"""

    def _blocks(self, analysis, review):
        """将 analysis + review 拆成通用块结构，供各格式渲染"""
        a = analysis or {}
        s = a.get('summary') or {}
        review = review or {}
        stats = _review_stats(review)
        return {
            'risks': a.get('risks') or [],
            'impact_areas': a.get('impact_areas') or [],
            'new_features': a.get('new_features') or [],
            'existing_features': a.get('existing_features') or [],
            'summary': s,
            'review': review or {},
            'stats': stats,
        }

    # ---- HTML ----
    def _to_html(self, b):
        def list_sec(title, items):
            if not items:
                return f'<details><summary><h3>{title}</h3></summary><p>无</p></details>'
            lis = ''.join(f'<li><b>{i.get("level","")}</b> {i.get("title","")} - {i.get("desc","")}</li>' if isinstance(i, dict) and 'title' in i else f'<li>{i}</li>' for i in items)
            return f'<details><summary><h3>{title}</h3></summary><ul>{lis}</ul></details>'
        h = ['<!DOCTYPE html><html lang="zh"><meta charset="utf-8">',
             '<style>body{font-family:微软雅黑,Arial,sans-serif;max-width:880px;margin:24px auto;line-height:1.7}'
             'h1{border-bottom:2px solid #6366f1;padding-bottom:8px}'
             'h3{color:#4f46e5;margin-top:18px}li{margin:4px 0}.badge{padding:1px 8px;border-radius:999px;color:#fff;font-size:12px}'
             '.high{background:#ef4444}.mid{background:#f59e0b}.low{background:#22c55e}'
             '.rv-score{font-size:20px;color:#4f46e5;font-weight:700}'
             'table{border-collapse:collapse;width:100%;margin-top:6px}td,th{border:1px solid #cbd5e1;padding:6px 8px;text-align:left;font-size:14px}'
             'th{background:#eef2ff}'
             'details{margin:14px 0;border:1px solid #e0e7ff;border-radius:8px;padding:8px 12px}'
             'summary{cursor:pointer;font-weight:700;color:#4f46e5}'
             'summary h2{display:inline;font-size:18px;margin:0;color:#4f46e5}'
             'summary h3{display:inline;font-size:16px;margin:0;color:#4f46e5}'
             'details[open]{border-color:#a5b4fc}'
             'details details{margin-top:8px;border-color:#eef2ff;background:#fbfcff}'
             'details details summary{font-size:14px}'
             'details details[open]{border-color:#c7d2fe}</style></head><body>',
             '<h1>需求评审报告</h1>']
        # 分析内容点
        h += ['<details><summary><h2>一、智能分析</h2></summary>',
              list_sec('风险点', b['risks']),
              list_sec('影响面', [{'level': x.get('level', ''), 'title': x.get('module', ''), 'desc': x.get('reason', '')} for x in b['impact_areas']])]
        h += ['<details><summary><h3>新增功能</h3></summary>' + ('<ul>' + ''.join(f'<li>{_esc(x)}</li>' for x in b['new_features']) + '</ul>' if b['new_features'] else '<p>无</p>') + '</details>']
        h += ['<details><summary><h3>既有功能</h3></summary>' + ('<ul>' + ''.join(f'<li>{_esc(x)}</li>' for x in b['existing_features']) + '</ul>' if b['existing_features'] else '<p>无</p>') + '</details>']
        s = b['summary']
        h += ['<details><summary><h3>报告摘要</h3></summary>',
              (f'<p><b>概述：</b>{_esc(s.get("overview"))}</p>' if s.get('overview') else ''),
              (f'<p><b>建议：</b>{_esc(s.get("recommendation"))}</p>' if s.get('recommendation') else ''),
              (f'<p><b>重点区域：</b>{_esc("、".join(s.get("focus_areas") or []))}</p>' if (s.get('focus_areas') or []) else ''),
              '</details>',
              '</details>']
        # 评审点
        r = b['review']
        st = b['stats']
        h += ['<details><summary><h2>二、用例评审点</h2></summary>',
              f'<p><span class="rv-score">{_esc(r.get("overall_score"))}</span> / 100 分 · '
              f'高<span class="badge high">{st["高"]}</span> 中<span class="badge mid">{st["中"]}</span> 低<span class="badge low">{st["低"]}</span></p>']
        if r.get('comment'):
            h.append(f'<p><b>总体意见：</b>{_esc(r["comment"])}</p>')
        if r.get('issues'):
            rows = ''.join(f'<tr><td>{_esc(i.get("case_index"))}</td><td><span class="badge {i.get("level","")}">{_esc(i.get("level"))}</span></td>'
                           f'<td>{_esc(i.get("problem"))}</td><td>{_esc(i.get("suggestion"))}</td></tr>' for i in r['issues'])
            h += ['<table><tr><th>用例序号</th><th>等级</th><th>问题</th><th>建议</th></tr>', rows, '</table>']
        else:
            h.append('<p>无评审问题</p>')
        h.append('</details>')
        h.append('</body></html>')
        return '\n'.join(h)

    # ---- Excel ----
    def _to_excel(self, b, path):
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        wb = Workbook()
        fill = PatternFill('solid', fgColor='4472C4')
        font = Font(color='FFFFFF', bold=True)
        ws = wb.active
        ws.title = '智能分析'
        ws.append(['风险点'])
        for c in ws[1]:
            c.fill, c.font, c.alignment = fill, font, Alignment(horizontal='center')
        ws.append(['等级', '标题', '描述'])
        for i in b['risks']:
            ws.append([i.get('level', ''), i.get('title', ''), i.get('desc', '')])
        ws = wb.create_sheet('影响面')
        ws.append(['等级', '模块', '原因'])
        for c in ws[1]:
            c.fill, c.font, c.alignment = fill, font, Alignment(horizontal='center')
        for i in b['impact_areas']:
            ws.append([i.get('level', ''), i.get('module', ''), i.get('reason', '')])
        ws = wb.create_sheet('功能')
        ws.append(['类别', '名称'])
        for c in ws[1]:
            c.fill, c.font, c.alignment = fill, font, Alignment(horizontal='center')
        for f in b['new_features']:
            ws.append(['新增', f])
        for f in b['existing_features']:
            ws.append(['既有', f])
        ws = wb.create_sheet('摘要')
        s = b['summary']
        ws.append(['概述', s.get('overview', '')])
        ws.append(['建议', s.get('recommendation', '')])
        ws.append(['重点区域', '、'.join(s.get('focus_areas') or [])])
        ws = wb.create_sheet('用例评审')
        r = b['review']
        st = b['stats']
        ws.append(['总分', r.get('overall_score', '')])
        ws.append(['高/中/低', f"{st['高']}/{st['中']}/{st['低']}"])
        ws.append(['总体意见', r.get('comment', '')])
        ws.append([])
        ws.append(['用例序号', '等级', '问题', '建议'])
        for c in ws[5]:
            c.fill, c.font, c.alignment = fill, font, Alignment(horizontal='center')
        for i in r.get('issues') or []:
            ws.append([i.get('case_index', ''), i.get('level', ''), i.get('problem', ''), i.get('suggestion', '')])
        wb.save(path)
        return path

    # ---- Word ----
    def _to_word(self, b, path):
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Pt
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        doc.add_heading('需求评审报告', level=1)
        doc.add_heading('一、智能分析', level=2)
        doc.add_heading('风险点', level=3)
        if b['risks']:
            for i in b['risks']:
                doc.add_paragraph(f"[{i.get('level','')}] {i.get('title','')}：{i.get('desc','')}")
        else:
            doc.add_paragraph('无')
        doc.add_heading('影响面', level=3)
        if b['impact_areas']:
            for i in b['impact_areas']:
                doc.add_paragraph(f"[{i.get('level','')}] {i.get('module','')}：{i.get('reason','')}")
        else:
            doc.add_paragraph('无')
        doc.add_heading('新增功能', level=3)
        doc.add_paragraph('、'.join(b['new_features']) if b['new_features'] else '无')
        doc.add_heading('既有功能', level=3)
        doc.add_paragraph('、'.join(b['existing_features']) if b['existing_features'] else '无')
        doc.add_heading('报告摘要', level=3)
        s = b['summary']
        if s.get('overview'):
            doc.add_paragraph(f'概述：{s.get("overview")}')
        if s.get('recommendation'):
            doc.add_paragraph(f'建议：{s.get("recommendation")}')
        if s.get('focus_areas'):
            doc.add_paragraph(f'重点区域：{"、".join(s.get("focus_areas"))}')
        doc.add_heading('二、用例评审点', level=2)
        r = b['review']
        st = b['stats']
        doc.add_paragraph(f'总分：{r.get("overall_score")} / 100   高{st["高"]} · 中{st["中"]} · 低{st["低"]}')
        if r.get('comment'):
            doc.add_paragraph(f'总体意见：{r.get("comment")}')
        if r.get('issues'):
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            for i, hd in enumerate(['用例序号', '等级', '问题', '建议']):
                table.rows[0].cells[i].text = hd
            for i in r['issues']:
                row = table.add_row().cells
                row[0].text = str(i.get('case_index', ''))
                row[1].text = str(i.get('level', ''))
                row[2].text = str(i.get('problem', ''))
                row[3].text = str(i.get('suggestion', ''))
        else:
            doc.add_paragraph('无评审问题')
        doc.save(path)
        return path

    # ---- XMind(.mm) ----
    def _to_xmind(self, b, path):
        out = ['<?xml version="1.0" encoding="UTF-8"?>', '<map version="1.0.1">', '<node TEXT="需求评审报告">']
        out.append('<node TEXT="一、智能分析">')
        out.append('<node TEXT="风险点">')
        for i in b['risks']:
            out.append(f'<node TEXT="[{_esc(i.get("level",""))}] {_esc(i.get("title",""))}：{_esc(i.get("desc",""))}"/>')
        out.append('</node>')
        out.append('<node TEXT="影响面">')
        for i in b['impact_areas']:
            out.append(f'<node TEXT="[{_esc(i.get("level",""))}] {_esc(i.get("module",""))}：{_esc(i.get("reason",""))}"/>')
        out.append('</node>')
        out.append(f'<node TEXT="新增功能">' + ''.join(f'<node TEXT="{_esc(x)}"/>' for x in b['new_features']) + '</node>')
        out.append(f'<node TEXT="既有功能">' + ''.join(f'<node TEXT="{_esc(x)}"/>' for x in b['existing_features']) + '</node>')
        s = b['summary']
        out.append(f'<node TEXT="报告摘要">{_esc(s.get("overview",""))}{"；" if s.get("overview") else ""}{_esc(s.get("recommendation",""))}</node>')
        out.append('</node>')
        r = b['review']
        st = b['stats']
        out.append(f'<node TEXT="二、用例评审点：{_esc(r.get("overall_score"))}/100 · 高{st["高"]}中{st["中"]}低{st["低"]}">')
        for i in r.get('issues') or []:
            out.append(f'<node TEXT="[{_esc(i.get("level",""))}] 用例{_esc(i.get("case_index"))} {_esc(i.get("problem",""))}（建议：{_esc(i.get("suggestion",""))}）"/>')
        out.append('</node>')
        out.append('</node>')
        out.append('</map>')
        with open(path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(out))
        return path

    def export(self, analysis, review, format_type, path):
        b = self._blocks(analysis, review)
        if format_type == 'html':
            with open(path, 'w', encoding='utf-8') as f:
                f.write(self._to_html(b))
        elif format_type == 'excel':
            return self._to_excel(b, path)
        elif format_type == 'word':
            return self._to_word(b, path)
        elif format_type == 'xmind':
            return self._to_xmind(b, path)
        else:
            raise ValueError(f'不支持的导出格式：{format_type}')
        return path
