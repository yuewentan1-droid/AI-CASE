"""操作手册生成 - 支持系统操作手册与用户操作手册导出

结合需求上下文（PRD/设计稿/所有输入）与用例（含优先级），生成通俗易懂、
图文并茂（表格 + 业务流程图）的操作手册。有 AI 处理器时走增强生成；
无 AI 时回退为基于用例分组的兜底结构。
"""
import xml.sax.saxutils as sax


def _esc(text):
    return sax.escape(str(text or ''))


class ManualGenerator:
    """基于用例与分析结果生成系统/用户操作手册（Word/HTML）"""

    def generate(self, cases, analysis, title, ai=None):
        """组装手册内容为结构化数据；ai 存在时走 AI 增强（概述/术语/模块说明/补优先级）"""
        if ai is not None:
            from com.aiase.services.export.manual_writer import ManualWriter
            return ManualWriter(ai).build(cases, analysis, title)
        # 兜底：无 AI 时保证每个用例有优先级，并按主模块分组
        mods = {}
        for c in cases:
            if not (c.get('优先级') or '').strip():
                c['优先级'] = '中'
            mods.setdefault(c.get('主模块', '其他'), []).append(c)
        return {
            'title': title,
            'overview': (analysis or {}).get('summary', {}).get('overview', ''),
            'glossary': [],
            'flowchart': (analysis or {}).get('flowchart', ''),
            'sections': [{'name': m, 'intro': '', 'cases': cs} for m, cs in mods.items()],
        }

    # ---- Word ----
    def export_word(self, content, path):
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        doc.add_heading(content['title'], level=0)
        # 一、操作概述
        doc.add_heading('一、操作概述', level=1)
        doc.add_paragraph(content['overview'] or '本手册依据需求文档与测试用例编写。')
        # 二、术语表
        if content.get('glossary'):
            doc.add_heading('二、术语表', level=1)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            for i, hd in enumerate(['术语', '通俗解释']):
                table.rows[0].cells[i].text = hd
            for g in content['glossary']:
                cells = table.add_row().cells
                cells[0].text = g.get('term', '')
                cells[1].text = g.get('desc', '')
        # 三、功能模块操作说明
        start = 3 if content.get('glossary') else 2
        doc.add_heading(f'{"三" if content.get("glossary") else "二"}、功能模块操作说明', level=1)
        for idx, sec in enumerate(content['sections'], 1):
            doc.add_heading(f'{start}.{idx} {sec["name"]}', level=2)
            if sec.get('intro'):
                doc.add_paragraph(sec['intro'])
            # 用例表格（表格并茂）
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            for i, hd in enumerate(['优先级', '用例标题', '操作步骤', '预期结果']):
                table.rows[0].cells[i].text = hd
            for c in sec['cases']:
                cells = table.add_row().cells
                cells[0].text = c.get('优先级', '中')
                cells[1].text = c.get('测试标题', '')
                self._cell_lines(cells[2], c.get('测试步骤', ''))
                self._cell_lines(cells[3], c.get('测试结果', ''))
        doc.save(path)
        return path

    @staticmethod
    def _cell_lines(cell, text):
        """将多行文本写入表格单元格（每行一个段落，保证换行可见）"""
        lines = [ln.strip() for ln in str(text or '').splitlines() if ln.strip()]
        if not lines:
            cell.text = ''
            return
        cell.paragraphs[0].text = lines[0]
        for ln in lines[1:]:
            cell.add_paragraph(ln)

    # ---- HTML ----
    def export_html(self, content, path):
        parts = [f'<h1>{_esc(content["title"])}</h1>']
        parts.append('<h2>一、操作概述</h2>')
        parts.append(f'<p>{_esc(content["overview"]) or "本手册依据需求文档与测试用例编写。"}</p>')
        # 业务流程图（图文并茂）：内联 mermaid，离线时显示源码
        if content.get('flowchart'):
            parts.append('<h2>二、业务流程总览</h2>')
            parts.append('<pre class="mermaid">' + _esc(content['flowchart']) + '</pre>')
        # 术语表
        if content.get('glossary'):
            gidx = 3
            parts.append('<h2>三、术语表</h2><table><tr><th>术语</th><th>通俗解释</th></tr>')
            for g in content['glossary']:
                parts.append(f'<tr><td>{_esc(g.get("term"))}</td><td>{_esc(g.get("desc"))}</td></tr>')
            parts.append('</table>')
        else:
            gidx = 2
        sec_h = '三' if content.get('glossary') else '二'
        parts.append(f'<h2>{sec_h}、功能模块操作说明</h2>')
        for idx, sec in enumerate(content['sections'], 1):
            parts.append(f'<details><summary><h3>{gidx}.{idx} {_esc(sec["name"])}</h3></summary>')
            if sec.get('intro'):
                parts.append(f'<p>{_esc(sec["intro"])}</p>')
            parts.append('<table><tr><th>优先级</th><th>用例标题</th><th>操作步骤</th><th>预期结果</th></tr>')
            for c in sec['cases']:
                steps = '<br>'.join(_esc(x) for x in str(c.get('测试步骤', '')).splitlines() if x.strip())
                results = '<br>'.join(_esc(x) for x in str(c.get('测试结果', '')).splitlines() if x.strip())
                parts.append(f'<tr><td>{_esc(c.get("优先级", "中"))}</td>'
                             f'<td>{_esc(c.get("测试标题", ""))}</td>'
                             f'<td>{steps}</td><td>{results}</td></tr>')
            parts.append('</table>')
            parts.append('</details>')
        html = f'''<!DOCTYPE html><html lang="zh"><head><meta charset="utf-8"><title>{_esc(content["title"])}</title>
<script type="module">import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs';mermaid.initialize({{startOnLoad:true}});</script>
<style>body{{font-family:Microsoft YaHei,微软雅黑,Arial,sans-serif;margin:30px;line-height:1.8;}}
h1{{border-bottom:2px solid #6366f1;padding-bottom:8px;}}h2{{color:#4f46e5;margin-top:24px;}}
h3{{color:#333;}}
details{{margin:12px 0;border:1px solid #e0e7ff;border-radius:8px;padding:8px 12px;}}
summary{{cursor:pointer;font-weight:700;color:#4f46e5;}}
summary h3{{display:inline;margin:0;font-size:16px;}}
details[open]{{border-color:#a5b4fc;}}
table{{border-collapse:collapse;width:100%;margin:10px 0;}}
td,th{{border:1px solid #cbd5e1;padding:6px 10px;text-align:left;font-size:14px;vertical-align:top;}}
th{{background:#eef2ff;}}</style></head>
<body>{''.join(parts)}</body></html>'''
        with open(path, 'w', encoding='utf-8') as f:
            f.write(html)
        return path
