"""Word导出"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from com.aiase.entity.template import resolve_fields
from com.aiase.services.export.base_exporter import BaseExporter


class WordExporter(BaseExporter):

    def export(self, cases, template, path):
        fields = resolve_fields(template)
        doc = Document()
        doc.add_heading('测试用例列表', level=1)
        # 设置中文字体
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        doc.add_paragraph(f'共 {len(cases)} 条用例')
        table = doc.add_table(rows=1, cols=len(fields))
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, f in enumerate(fields):
            hdr[i].text = f
        for c in cases:
            row = table.add_row().cells
            for i, f in enumerate(fields):
                self._cell_lines(row[i], c.get(f, ''))
        doc.save(path)
        return path

    @staticmethod
    def _cell_lines(cell, value):
        """将多行文本写入单元格（每行一个段落），保证步骤/结果换行可见、一一对应"""
        lines = [ln for ln in str(value if value is not None else '').split('\n')]
        if not lines:
            cell.text = ''
            return
        cell.text = lines[0]
        for ln in lines[1:]:
            cell.add_paragraph(ln)
