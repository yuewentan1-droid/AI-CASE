"""txt、md、word文档处理"""
from com.aiase.services.input.base_handler import BaseHandler
from com.aiase.utils.file_utils import read_text_file


class TextHandler(BaseHandler):
    """处理 .txt 与 .md 纯文本文件"""

    def parse(self, file_path):
        return read_text_file(file_path)


class WordHandler(BaseHandler):
    """处理 .docx 文档，提取段落与表格"""

    def parse(self, file_path):
        from docx import Document
        doc = Document(file_path)
        parts = ['# 文档内容']
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            parts.append('--- 表格 ---')
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                parts.append(' | '.join(cells))
        return '\n'.join(parts)
